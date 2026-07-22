import os
import io
import json
import shutil
import tempfile
import numpy as np
from pathlib import Path
from dotenv import load_dotenv
from fastapi import FastAPI, UploadFile, File, HTTPException, Header, Query
from fastapi.responses import HTMLResponse, FileResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from typing import Optional

from analyzer import process_excel, build_bodega_prompt, search_inventory, get_vendors, get_vendor_items


def to_serializable(obj):
    if isinstance(obj, dict):
        return {k: to_serializable(v) for k, v in obj.items()}
    if isinstance(obj, (list, tuple)):
        return [to_serializable(v) for v in obj]
    if isinstance(obj, (np.integer,)):
        return int(obj)
    if isinstance(obj, (np.floating,)):
        return float(obj)
    if isinstance(obj, np.ndarray):
        return obj.tolist()
    return obj


GEMINI_MODELS = ["gemini-3.1-flash-lite", "gemini-3-flash-preview", "gemini-3.5-flash"]
CLAUDE_MODELS = ["claude-sonnet-4-20250514", "claude-3-5-haiku-20241022"]
OPENAI_MODELS = ["gpt-4o", "gpt-4o-mini"]


def _call_gemini(prompt: str, api_key: str) -> str:
    import google.generativeai as genai
    genai.configure(api_key=api_key)
    last_error = None
    for model_name in GEMINI_MODELS:
        try:
            model = genai.GenerativeModel(model_name)
            response = model.generate_content(prompt)
            return response.text
        except Exception as e:
            last_error = e
            err_str = str(e).lower()
            if "429" in err_str or "quota" in err_str or "rate" in err_str:
                continue
            raise e
    raise Exception(f"Todos los modelos Gemini fallaron. Ultimo error: {last_error}")


def _call_claude(prompt: str, api_key: str) -> str:
    import time
    import anthropic
    client = anthropic.Anthropic(api_key=api_key)
    last_error = None
    for model_name in CLAUDE_MODELS:
        for attempt in range(3):
            try:
                response = client.messages.create(
                    model=model_name,
                    max_tokens=8192,
                    messages=[{"role": "user", "content": prompt}],
                )
                return response.content[0].text
            except Exception as e:
                last_error = e
                err_str = str(e).lower()
                if "429" in err_str or "rate" in err_str or "overloaded" in err_str:
                    if attempt < 2:
                        time.sleep(5 * (attempt + 1))
                        continue
                    else:
                        break
                else:
                    raise e
    raise Exception(f"Todos los modelos Claude fallaron. Ultimo error: {last_error}")


def _call_openai(prompt: str, api_key: str) -> str:
    import time
    from openai import OpenAI
    client = OpenAI(api_key=api_key)
    last_error = None
    for model_name in OPENAI_MODELS:
        for attempt in range(3):
            try:
                response = client.chat.completions.create(
                    model=model_name,
                    messages=[{"role": "user", "content": prompt}],
                    max_tokens=8192,
                )
                return response.choices[0].message.content
            except Exception as e:
                last_error = e
                err_str = str(e).lower()
                if "429" in err_str or "rate" in err_str or "quota" in err_str:
                    if attempt < 2:
                        time.sleep(5 * (attempt + 1))
                        continue
                    else:
                        break
                else:
                    raise e
    raise Exception(f"Todos los modelos OpenAI fallaron. Ultimo error: {last_error}")


def _call_ai(prompt: str, api_key: str, provider: str = "gemini") -> str:
    if provider == "claude":
        return _call_claude(prompt, api_key)
    elif provider == "openai":
        return _call_openai(prompt, api_key)
    else:
        return _call_gemini(prompt, api_key)


CHART_COLORS = ['#4F46E5', '#7C3AED', '#2563EB', '#3B82F6', '#8B5CF6',
                '#A78BFA', '#C4B5FD', '#DDD6FE', '#6366F1', '#818CF8']


def _chart_valor_por_bodega(bodega_data: list[dict]) -> io.BytesIO:
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt

    bodegas = bodega_data[:10]
    if not bodegas:
        return None

    nombres = [b['bodega'][:30] for b in reversed(bodegas)]
    valores = [b['valor_total'] for b in reversed(bodegas)]
    comprometido = [b['cant_comprometida'] for b in reversed(bodegas)]

    fig, ax = plt.subplots(figsize=(9, max(3.5, len(nombres) * 0.55)))
    y_pos = range(len(nombres))

    bars1 = ax.barh(y_pos, valores, height=0.35, label='Valor Total',
                     color=CHART_COLORS[0], alpha=0.85, edgecolor='white', linewidth=0.5)
    ax.set_yticks(y_pos)
    ax.set_yticklabels(nombres, fontsize=8)
    ax.set_xlabel('Valor ($)', fontsize=9)
    ax.set_title('Top Bodegas por Valor de Inventario', fontsize=12, fontweight='bold',
                  color='#1E293B', pad=12)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.grid(axis='x', alpha=0.2, linestyle='--')
    ax.tick_params(axis='both', labelsize=8)

    for bar in bars1:
        width = bar.get_width()
        if width > 0:
            ax.text(width + max(valores) * 0.01, bar.get_y() + bar.get_height()/2,
                     f'${width:,.0f}', va='center', fontsize=7, color='#475569')

    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    plt.close(fig)
    buf.seek(0)
    return buf


def _chart_estado_inventario(bodega_data: dict) -> io.BytesIO:
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt

    estados = bodega_data.get('por_estado', {})
    if not estados:
        return None

    labels = list(estados.keys())
    sizes = [estados[e]['existencia'] for e in labels]
    colors_map = {
        'E1 - LIQUIDACION': '#F87171', 'E2 - EVENTO': '#FBBF24', 'E3 - MTO': '#34D399',
        'E4 - STOCK': '#60A5FA', 'E5 - KOOK': '#A78BFA'
    }
    colors = [colors_map.get(l, '#94A3B8') for l in labels]

    fig, ax = plt.subplots(figsize=(6, 4))
    wedges, texts, autotexts = ax.pie(
        sizes, labels=labels, autopct='%1.0f%%', startangle=90,
        colors=colors, pctdistance=0.75, labeldistance=1.15,
        wedgeprops=dict(width=0.5, edgecolor='white', linewidth=2)
    )
    for t in texts:
        t.set_fontsize(8)
    for t in autotexts:
        t.set_fontsize(7)
        t.set_fontweight('bold')

    ax.set_title('Distribucion por Estado de Inventario', fontsize=11,
                  fontweight='bold', color='#1E293B', pad=12)
    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    plt.close(fig)
    buf.seek(0)
    return buf


def _chart_lineas(bodega_data: dict) -> io.BytesIO:
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt

    lineas = bodega_data.get('por_linea', {})
    if not lineas:
        return None

    items = sorted(lineas.items(), key=lambda x: x[1]['valor_total'], reverse=True)[:8]
    nombres = [l[0][:25] for l in items]
    valores = [l[1]['valor_total'] for l in items]
    existencias = [l[1]['existencia'] for l in items]

    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(10, 4))

    bars1 = ax1.bar(range(len(nombres)), existencias, color=CHART_COLORS[0], alpha=0.85,
                     edgecolor='white', linewidth=0.5)
    ax1.set_xticks(range(len(nombres)))
    ax1.set_xticklabels(nombres, rotation=35, ha='right', fontsize=7)
    ax1.set_ylabel('Existencia', fontsize=8)
    ax1.set_title('Por Existencia', fontsize=10, fontweight='bold', color='#1E293B')
    ax1.spines['top'].set_visible(False)
    ax1.spines['right'].set_visible(False)
    ax1.grid(axis='y', alpha=0.2, linestyle='--')
    ax1.tick_params(axis='both', labelsize=7)
    for bar in bars1:
        h = bar.get_height()
        ax1.text(bar.get_x() + bar.get_width()/2, h + max(existencias)*0.01,
                  f'{h:,.0f}', ha='center', va='bottom', fontsize=6, color='#475569')

    bars2 = ax2.bar(range(len(nombres)), valores, color=CHART_COLORS[2], alpha=0.85,
                     edgecolor='white', linewidth=0.5)
    ax2.set_xticks(range(len(nombres)))
    ax2.set_xticklabels(nombres, rotation=35, ha='right', fontsize=7)
    ax2.set_ylabel('Valor ($)', fontsize=8)
    ax2.set_title('Por Valor', fontsize=10, fontweight='bold', color='#1E293B')
    ax2.spines['top'].set_visible(False)
    ax2.spines['right'].set_visible(False)
    ax2.grid(axis='y', alpha=0.2, linestyle='--')
    ax2.tick_params(axis='both', labelsize=7)
    for bar in bars2:
        h = bar.get_height()
        ax2.text(bar.get_x() + bar.get_width()/2, h + max(valores)*0.01,
                  f'${h:,.0f}', ha='center', va='bottom', fontsize=6, color='#475569')

    plt.suptitle('Lineas de Producto', fontsize=12, fontweight='bold',
                  color='#1E293B', y=1.02)
    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    plt.close(fig)
    buf.seek(0)
    return buf


def _verify_ai_key(api_key: str, provider: str = "gemini") -> dict:
    models_tried = []
    try:
        if provider == "claude":
            import anthropic
            client = anthropic.Anthropic(api_key=api_key)
            response = client.messages.create(
                model=CLAUDE_MODELS[0],
                max_tokens=10,
                messages=[{"role": "user", "content": "Responde solo: OK"}],
            )
            return {"success": True, "model": CLAUDE_MODELS[0], "message": "Conexion exitosa con " + CLAUDE_MODELS[0]}
        elif provider == "openai":
            from openai import OpenAI
            client = OpenAI(api_key=api_key)
            response = client.chat.completions.create(
                model=OPENAI_MODELS[0],
                messages=[{"role": "user", "content": "Responde solo: OK"}],
                max_tokens=10,
            )
            return {"success": True, "model": OPENAI_MODELS[0], "message": "Conexion exitosa con " + OPENAI_MODELS[0]}
        else:
            import google.generativeai as genai
            genai.configure(api_key=api_key)
            last_err = ""
            for model_name in GEMINI_MODELS:
                models_tried.append(model_name)
                try:
                    model = genai.GenerativeModel(model_name)
                    response = model.generate_content("Responde solo: OK")
                    return {"success": True, "model": model_name, "message": "Conexion exitosa con " + model_name}
                except Exception as e:
                    last_err = str(e)
                    continue
            raise Exception(last_err or "Todos los modelos fallaron")
    except Exception as e:
        err_str = str(e).lower()
        if "429" in err_str or "quota" in err_str or "rate" in err_str:
            tried = ", ".join(models_tried) if models_tried else provider
            raise HTTPException(status_code=429, detail=f"Cuota agotada. Modelos intentados: {tried}. Espera o usa otra API key.")
        if "invalid" in err_str or "unauthorized" in err_str or "401" in err_str:
            raise HTTPException(status_code=400, detail="API key invalida para " + provider)
        raise HTTPException(status_code=400, detail=f"Error verificando {provider}: {str(e)}")

load_dotenv()

app = FastAPI(title="Analisis de Inventario", version="1.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)

current_data = {}


@app.get("/", response_class=HTMLResponse)
async def serve_index():
    html_path = Path("index.html")
    if html_path.exists():
        return HTMLResponse(content=html_path.read_text(encoding="utf-8"))
    return HTMLResponse(content="<h1>index.html no encontrado</h1>", status_code=404)


@app.post("/verify-key")
async def verify_key(
    x_api_key: Optional[str] = Header(None),
    x_provider: Optional[str] = Header(None),
):
    api_key = x_api_key or ""
    provider = x_provider or "gemini"
    if not api_key:
        raise HTTPException(status_code=400, detail="No se proporciono API key.")
    return _verify_ai_key(api_key, provider)


@app.post("/upload")
async def upload_excel(file: UploadFile = File(...)):
    if not file.filename:
        raise HTTPException(status_code=400, detail="No se envio ningun archivo. Selecciona un archivo Excel.")

    if not file.filename.endswith((".xlsx", ".xls")):
        raise HTTPException(status_code=400, detail="Formato no valido. Solo se permiten archivos Excel (.xlsx)")

    filepath = UPLOAD_DIR / file.filename
    try:
        with open(filepath, "wb") as f:
            shutil.copyfileobj(file.file, f)
    except PermissionError:
        raise HTTPException(status_code=500, detail="No se pudo guardar el archivo. Verifica que no este abierto en Excel y que tengas permisos de escritura.")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error al guardar el archivo: {str(e)}")

    if filepath.stat().st_size == 0:
        filepath.unlink(missing_ok=True)
        raise HTTPException(status_code=400, detail="El archivo esta vacio. Selecciona un archivo Excel valido.")

    try:
        result = process_excel(str(filepath))
        current_data["result"] = to_serializable(result)
        current_data["filename"] = file.filename
        current_data["filepath"] = str(filepath)

        import pandas as pd
        try:
            raw_df = pd.read_excel(str(filepath), engine='calamine')
        except Exception:
            try:
                raw_df = pd.read_excel(str(filepath), engine='openpyxl')
            except Exception:
                raw_df = None
        current_data["raw_df"] = raw_df
    except ValueError as e:
        filepath.unlink(missing_ok=True)
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        filepath.unlink(missing_ok=True)
        raise HTTPException(status_code=500, detail=f"Error al procesar el archivo: {str(e)}")

    return {
        "success": True,
        "filename": file.filename,
        "total_rows": result["total_raw_rows"],
        "bodegas": len(result["bodega_metrics"]),
        "bodega_filter": result["bodega_filter"],
        "detected_columns": result["detected_columns"],
    }


@app.get("/data")
async def get_data():
    if not current_data.get("result"):
        raise HTTPException(status_code=400, detail="No hay datos cargados. Sube un Excel primero.")
    return current_data["result"]


@app.get("/search")
async def search(q: str = Query(""), limit: int = Query(50)):
    if not current_data.get("raw_df") is not None:
        raise HTTPException(status_code=400, detail="No hay datos cargados. Sube un Excel primero.")
    if not q or len(q.strip()) < 2:
        return {"results": [], "query": q}
    results = search_inventory(q, current_data["raw_df"], limit=limit)
    return {"results": results, "query": q}


@app.post("/refilter")
async def refilter(
    canal: Optional[str] = Query(None),
    categoria: Optional[str] = Query(None),
    estado: Optional[str] = Query(None),
    linea: Optional[str] = Query(None),
):
    if not current_data.get("filepath"):
        raise HTTPException(status_code=400, detail="No hay datos cargados. Sube un Excel primero.")

    try:
        result = process_excel(
            current_data["filepath"],
            canal=canal or None,
            categoria=categoria or None,
            estado=estado or None,
            linea=linea or None,
        )
        current_data["result"] = to_serializable(result)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error al re-filtrar: {str(e)}")

    return current_data["result"]


@app.post("/analyze/{bodega_name}")
async def analyze_bodega(
    bodega_name: str,
    x_api_key: Optional[str] = Header(None),
    x_provider: Optional[str] = Header(None),
):
    if not current_data.get("result"):
        raise HTTPException(status_code=400, detail="No hay datos cargados. Sube un Excel primero.")

    api_key = x_api_key or os.getenv("GEMINI_API_KEY", "")
    provider = x_provider or "gemini"
    if not api_key:
        raise HTTPException(status_code=400, detail="No se proporciono API key.")

    result = current_data["result"]
    bodega_metrics = result["bodega_metrics"]
    global_summary = result["global_summary"]

    bodega_data = None
    for m in bodega_metrics:
        if m["bodega"].strip().upper() == bodega_name.strip().upper():
            bodega_data = m
            break

    if not bodega_data:
        raise HTTPException(status_code=404, detail=f"Bodega '{bodega_name}' no encontrada.")

    bodega_filter = result.get("bodega_filter", "")
    prompt = build_bodega_prompt(bodega_data, global_summary, bodega_filter)

    try:
        informe = _call_ai(prompt, api_key, provider)
    except Exception as e:
        err_str = str(e).lower()
        if "429" in err_str or "quota" in err_str or "rate" in err_str:
            raise HTTPException(status_code=429, detail="Cuota agotada. Verifica tu plan o usa otra API key.")
        raise HTTPException(status_code=500, detail=f"Error con {provider}: {str(e)}")

    cache_key = f"report_{bodega_name.strip().upper()}"
    current_data[cache_key] = informe

    return {
        "bodega": bodega_name,
        "informe": informe,
        "metricas": bodega_data,
    }


@app.post("/analyze-all")
async def analyze_all(
    x_api_key: Optional[str] = Header(None),
    x_provider: Optional[str] = Header(None),
):
    if not current_data.get("result"):
        raise HTTPException(status_code=400, detail="No hay datos cargados. Sube un Excel primero.")

    api_key = x_api_key or os.getenv("GEMINI_API_KEY", "")
    provider = x_provider or "gemini"
    if not api_key:
        raise HTTPException(status_code=400, detail="No se proporciono API key.")

    result = current_data["result"]
    bodega_metrics = result["bodega_metrics"]
    global_summary = result["global_summary"]
    bodega_filter = result.get("bodega_filter", "")

    informes = []
    for m in bodega_metrics:
        try:
            prompt = build_bodega_prompt(m, global_summary, bodega_filter)
            informe_text = _call_ai(prompt, api_key, provider)
            informes.append({
                "bodega": m["bodega"],
                "informe": informe_text,
                "metricas": m,
            })
        except Exception as e:
            informes.append({
                "bodega": m["bodega"],
                "informe": f"Error con {provider}: {str(e)}",
                "metricas": m,
            })

    return {"informes": informes}


@app.get("/export/{bodega_name}")
async def export_bodega(bodega_name: str):
    if not current_data.get("result"):
        raise HTTPException(status_code=400, detail="No hay datos cargados.")

    result = current_data["result"]
    bodega_data = None
    for m in result["bodega_metrics"]:
        if m["bodega"].strip().upper() == bodega_name.strip().upper():
            bodega_data = m
            break

    if not bodega_data:
        raise HTTPException(status_code=404, detail=f"Bodega '{bodega_name}' no encontrada.")

    return bodega_data


def _get_bodega_data(bodega_name: str) -> dict:
    if not current_data.get("result"):
        raise HTTPException(status_code=400, detail="No hay datos cargados.")
    result = current_data["result"]
    for m in result["bodega_metrics"]:
        if m["bodega"].strip().upper() == bodega_name.strip().upper():
            return m
    raise HTTPException(status_code=404, detail=f"Bodega '{bodega_name}' no encontrada.")


def _generate_bodega_report(bodega_name: str, api_key: str, provider: str = "gemini") -> str:
    result = current_data["result"]
    bodega_data = _get_bodega_data(bodega_name)
    global_summary = result["global_summary"]

    cache_key = f"report_{bodega_name.strip().upper()}"
    if cache_key in current_data:
        return current_data[cache_key]

    bodega_filter = result.get("bodega_filter", "")
    prompt = build_bodega_prompt(bodega_data, global_summary, bodega_filter)
    informe = _call_ai(prompt, api_key, provider)
    current_data[cache_key] = informe
    return informe


@app.post("/generate-and-cache/{bodega_name}")
async def generate_and_cache(
    bodega_name: str,
    x_api_key: Optional[str] = Header(None),
    x_provider: Optional[str] = Header(None),
):
    api_key = x_api_key or os.getenv("GEMINI_API_KEY", "")
    provider = x_provider or "gemini"
    if not api_key:
        raise HTTPException(status_code=400, detail="No se proporciono API key.")

    try:
        informe = _generate_bodega_report(bodega_name, api_key, provider)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error al generar informe: {str(e)}")

    return {"bodega": bodega_name, "informe": informe}


@app.get("/download/{bodega_name}/word")
async def download_word(bodega_name: str):
    bodega_data = _get_bodega_data(bodega_name)
    result = current_data["result"]
    cache_key = f"report_{bodega_name.strip().upper()}"
    informe = current_data.get(cache_key, "")

    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    title = doc.add_heading('', level=0)
    run = title.add_run(f'Informe de Inventario - {bodega_name}')
    run.font.color.rgb = RGBColor(79, 70, 229)
    run.font.size = Pt(22)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Generado con IA')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph()

    doc.add_heading('Metricas de la Bodega', level=1)
    table = doc.add_table(rows=10, cols=2, style='Light Shading Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    metrics_data = [
        ('Existencia Total', f"{bodega_data['existencia']:,.0f}"),
        ('Cant. Comprometida', f"{bodega_data['cant_comprometida']:,.0f}"),
        ('Cant. Disponible', f"{bodega_data['cant_disponible']:,.0f}"),
        ('% Compromiso', f"{bodega_data['compromiso_pct']}%"),
        ('Valor Total', f"${bodega_data['valor_total']:,.0f}"),
        ('Costo Prom. Total', f"${bodega_data['costo_prom_total']:,.0f}"),
        ('Precio Unitario Prom.', f"${bodega_data['precio_unitario']:,.0f}"),
        ('Margen Promedio', f"{bodega_data['margen_promedio']:.1f}%"),
        ('Referencias Unicas', str(bodega_data['refs_unicas'])),
        ('Total Registros', str(bodega_data['total_registros'])),
    ]

    for i, (label, value) in enumerate(metrics_data):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    if bodega_data.get('por_estado'):
        doc.add_paragraph()
        doc.add_heading('Por Estado', level=2)
        estado_table = doc.add_table(rows=len(bodega_data['por_estado']) + 1, cols=4, style='Light List Accent 1')
        header_row = estado_table.rows[0]
        header_row.cells[0].text = 'Estado'
        header_row.cells[1].text = 'Existencia'
        header_row.cells[2].text = 'Comprometida'
        header_row.cells[3].text = 'Valor'
        for cell_idx in range(4):
            for run in header_row.cells[cell_idx].paragraphs[0].runs:
                run.bold = True
        for i, (estado, data) in enumerate(bodega_data['por_estado'].items()):
            row = estado_table.rows[i + 1]
            row.cells[0].text = estado
            row.cells[1].text = f"{data['existencia']:,.0f}"
            row.cells[2].text = f"{data['comprometida']:,.0f}"
            row.cells[3].text = f"${data['valor_total']:,.0f}"

        chart_buf = _chart_estado_inventario(bodega_data)
        if chart_buf:
            doc.add_paragraph()
            doc.add_picture(chart_buf, width=Inches(5.0))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if bodega_data.get('por_linea'):
        doc.add_paragraph()
        doc.add_heading('Lineas de Producto', level=2)
        linea_table = doc.add_table(rows=len(bodega_data['por_linea']) + 1, cols=5, style='Light List Accent 1')
        header_row = linea_table.rows[0]
        for idx, h in enumerate(['Linea', 'Existencia', 'Comprometida', 'Valor', 'Margen']):
            header_row.cells[idx].text = h
            for run in header_row.cells[idx].paragraphs[0].runs:
                run.bold = True
        for i, (linea, data) in enumerate(bodega_data['por_linea'].items()):
            row = linea_table.rows[i + 1]
            row.cells[0].text = linea
            row.cells[1].text = f"{data['existencia']:,.0f}"
            row.cells[2].text = f"{data['comprometida']:,.0f}"
            row.cells[3].text = f"${data['valor_total']:,.0f}"
            row.cells[4].text = f"{data['margen']:.1f}%"

        chart_buf = _chart_lineas(bodega_data)
        if chart_buf:
            doc.add_paragraph()
            doc.add_picture(chart_buf, width=Inches(5.8))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if bodega_data.get('por_proveedor'):
        doc.add_paragraph()
        doc.add_heading('Proveedores', level=2)
        prov_table = doc.add_table(rows=len(bodega_data['por_proveedor']) + 1, cols=5, style='Light List Accent 1')
        header_row = prov_table.rows[0]
        for idx, h in enumerate(['Proveedor', 'Existencia', 'Comprometida', 'Valor', 'Margen']):
            header_row.cells[idx].text = h
            for run in header_row.cells[idx].paragraphs[0].runs:
                run.bold = True
        for i, (prov, data) in enumerate(bodega_data['por_proveedor'].items()):
            row = prov_table.rows[i + 1]
            row.cells[0].text = prov
            row.cells[1].text = f"{data['existencia']:,.0f}"
            row.cells[2].text = f"{data['comprometida']:,.0f}"
            row.cells[3].text = f"${data['valor_total']:,.0f}"
            row.cells[4].text = f"{data['margen']:.1f}%"

    if bodega_data.get('top_items'):
        doc.add_paragraph()
        doc.add_heading('Top Items por Valor', level=2)
        items_table = doc.add_table(rows=len(bodega_data['top_items']) + 1, cols=5, style='Light List Accent 1')
        header_row = items_table.rows[0]
        for idx, h in enumerate(['Item', 'Existencia', 'Comprometida', 'Valor', 'Precio']):
            header_row.cells[idx].text = h
            for run in header_row.cells[idx].paragraphs[0].runs:
                run.bold = True
        for i, item in enumerate(bodega_data['top_items']):
            row = items_table.rows[i + 1]
            row.cells[0].text = item['item']
            row.cells[1].text = f"{item['existencia']:,.0f}"
            row.cells[2].text = f"{item['comprometida']:,.0f}"
            row.cells[3].text = f"${item['valor_total']:,.0f}"
            row.cells[4].text = f"${item['precio_unitario']:,.0f}"

    if informe:
        doc.add_page_break()
        doc.add_heading('Informe de IA', level=1)

        lines = informe.split('\n')
        i = 0
        while i < len(lines):
            stripped = lines[i].strip()

            if stripped.startswith('### '):
                doc.add_heading(stripped[4:], level=2)
            elif stripped.startswith('## '):
                doc.add_heading(stripped[3:], level=1)
            elif stripped.startswith('# '):
                doc.add_heading(stripped[2:], level=0)

            elif stripped.startswith('|') and '|' in stripped[1:]:
                table_rows = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    row_line = lines[i].strip()
                    cells = [c.strip() for c in row_line.split('|')[1:-1]]
                    is_sep = all(c.replace('-', '').replace(':', '').strip() == '' for c in cells)
                    if not is_sep:
                        table_rows.append(cells)
                    i += 1
                i -= 1

                if table_rows:
                    num_cols = max(len(r) for r in table_rows)
                    t = doc.add_table(rows=len(table_rows), cols=num_cols, style='Light Shading Accent 1')
                    t.alignment = WD_TABLE_ALIGNMENT.CENTER
                    for ri, row_data in enumerate(table_rows):
                        for ci, cell_text in enumerate(row_data):
                            if ci < num_cols:
                                t.rows[ri].cells[ci].text = cell_text
                                for paragraph in t.rows[ri].cells[ci].paragraphs:
                                    for run in paragraph.runs:
                                        run.font.size = Pt(10)
                                        if ri == 0:
                                            run.bold = True

            elif stripped.startswith('- '):
                doc.add_paragraph(stripped[2:], style='List Bullet')
            elif stripped:
                p = doc.add_paragraph()
                parts = stripped.split('**')
                for j, part in enumerate(parts):
                    run = p.add_run(part)
                    run.font.size = Pt(11)
                    if j % 2 == 1:
                        run.bold = True

            i += 1

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    safe_name = bodega_name.replace(' ', '_').replace('.', '').replace('/', '_')
    return StreamingResponse(
        buffer,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f'attachment; filename="Informe_{safe_name}.docx"'}
    )


@app.get("/download/{bodega_name}/excel")
async def download_excel(bodega_name: str):
    bodega_data = _get_bodega_data(bodega_name)
    result = current_data["result"]

    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    wb = openpyxl.Workbook()

    header_font = Font(name='Calibri', bold=True, size=12, color='FFFFFF')
    header_fill = PatternFill(start_color='4F46E5', end_color='4F46E5', fill_type='solid')
    title_font = Font(name='Calibri', bold=True, size=16, color='4F46E5')
    subtitle_font = Font(name='Calibri', bold=True, size=10, color='64748B')
    metric_label_font = Font(name='Calibri', bold=True, size=11)
    metric_value_font = Font(name='Calibri', size=11)
    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )

    ws = wb.active
    ws.title = 'Metricas Bodega'

    ws.merge_cells('A1:D1')
    ws['A1'] = f'Informe de Inventario - {bodega_name}'
    ws['A1'].font = title_font
    ws['A1'].alignment = Alignment(horizontal='center')

    ws['A4'] = 'Metrica'
    ws['B4'] = 'Valor'
    ws['A4'].font = header_font
    ws['A4'].fill = header_fill
    ws['B4'].font = header_font
    ws['B4'].fill = header_fill
    ws['A4'].border = thin_border
    ws['B4'].border = thin_border

    metrics_data = [
        ('Existencia Total', bodega_data['existencia']),
        ('Cant. Comprometida', bodega_data['cant_comprometida']),
        ('Cant. Disponible', bodega_data['cant_disponible']),
        ('% Compromiso', bodega_data['compromiso_pct']),
        ('Valor Total', bodega_data['valor_total']),
        ('Costo Prom. Total', bodega_data['costo_prom_total']),
        ('Precio Unitario Prom.', bodega_data['precio_unitario']),
        ('Margen Promedio (%)', bodega_data['margen_promedio']),
        ('Referencias Unicas', bodega_data['refs_unicas']),
        ('Items Sin Movimiento', bodega_data['sin_movimiento']),
        ('Total Registros', bodega_data['total_registros']),
    ]

    for i, (label, value) in enumerate(metrics_data):
        row = i + 5
        ws.cell(row=row, column=1, value=label).font = metric_label_font
        cell = ws.cell(row=row, column=2, value=value)
        cell.font = metric_value_font
        cell.number_format = '#,##0'
        ws.cell(row=row, column=1).border = thin_border
        ws.cell(row=row, column=2).border = thin_border

    ws.column_dimensions['A'].width = 30
    ws.column_dimensions['B'].width = 20

    if bodega_data.get('por_estado'):
        ws2 = wb.create_sheet('Por Estado')
        headers2 = ['Estado', 'Existencia', 'Comprometida', 'Valor', 'Registros']
        for j, h in enumerate(headers2):
            cell = ws2.cell(row=1, column=j + 1, value=h)
            cell.font = header_font
            cell.fill = header_fill
        for i, (estado, data) in enumerate(bodega_data['por_estado'].items()):
            ws2.cell(row=i + 2, column=1, value=estado)
            ws2.cell(row=i + 2, column=2, value=data['existencia'])
            ws2.cell(row=i + 2, column=3, value=data['comprometida'])
            ws2.cell(row=i + 2, column=4, value=data['valor_total'])
            ws2.cell(row=i + 2, column=5, value=data['registros'])
        ws2.column_dimensions['A'].width = 25
        ws2.column_dimensions['B'].width = 15
        ws2.column_dimensions['C'].width = 15
        ws2.column_dimensions['D'].width = 20
        ws2.column_dimensions['E'].width = 12

    if bodega_data.get('por_linea'):
        ws3 = wb.create_sheet('Por Linea')
        headers3 = ['Linea', 'Existencia', 'Comprometida', 'Disponible', 'Valor', 'Margen', 'Registros']
        for j, h in enumerate(headers3):
            cell = ws3.cell(row=1, column=j + 1, value=h)
            cell.font = header_font
            cell.fill = header_fill
        for i, (linea, data) in enumerate(bodega_data['por_linea'].items()):
            ws3.cell(row=i + 2, column=1, value=linea)
            ws3.cell(row=i + 2, column=2, value=data['existencia'])
            ws3.cell(row=i + 2, column=3, value=data['comprometida'])
            ws3.cell(row=i + 2, column=4, value=data['disponible'])
            ws3.cell(row=i + 2, column=5, value=data['valor_total'])
            ws3.cell(row=i + 2, column=6, value=data['margen'])
            ws3.cell(row=i + 2, column=7, value=data['registros'])
        ws3.column_dimensions['A'].width = 35
        ws3.column_dimensions['B'].width = 15
        ws3.column_dimensions['C'].width = 15
        ws3.column_dimensions['D'].width = 15
        ws3.column_dimensions['E'].width = 20
        ws3.column_dimensions['F'].width = 10
        ws3.column_dimensions['G'].width = 12

    if bodega_data.get('por_categoria'):
        ws4 = wb.create_sheet('Por Categoria')
        headers4 = ['Categoria', 'Existencia', 'Comprometida', 'Valor', 'Margen', 'Registros']
        for j, h in enumerate(headers4):
            cell = ws4.cell(row=1, column=j + 1, value=h)
            cell.font = header_font
            cell.fill = header_fill
        for i, (cat, data) in enumerate(bodega_data['por_categoria'].items()):
            ws4.cell(row=i + 2, column=1, value=cat)
            ws4.cell(row=i + 2, column=2, value=data['existencia'])
            ws4.cell(row=i + 2, column=3, value=data['comprometida'])
            ws4.cell(row=i + 2, column=4, value=data['valor_total'])
            ws4.cell(row=i + 2, column=5, value=data['margen'])
            ws4.cell(row=i + 2, column=6, value=data['registros'])
        ws4.column_dimensions['A'].width = 30
        ws4.column_dimensions['B'].width = 15
        ws4.column_dimensions['C'].width = 15
        ws4.column_dimensions['D'].width = 20
        ws4.column_dimensions['E'].width = 10
        ws4.column_dimensions['F'].width = 12

    if bodega_data.get('por_canal'):
        ws5 = wb.create_sheet('Por Canal')
        headers5 = ['Canal', 'Existencia', 'Comprometida', 'Valor', 'Registros']
        for j, h in enumerate(headers5):
            cell = ws5.cell(row=1, column=j + 1, value=h)
            cell.font = header_font
            cell.fill = header_fill
        for i, (canal, data) in enumerate(bodega_data['por_canal'].items()):
            ws5.cell(row=i + 2, column=1, value=canal)
            ws5.cell(row=i + 2, column=2, value=data['existencia'])
            ws5.cell(row=i + 2, column=3, value=data['comprometida'])
            ws5.cell(row=i + 2, column=4, value=data['valor_total'])
            ws5.cell(row=i + 2, column=5, value=data['registros'])
        ws5.column_dimensions['A'].width = 35
        ws5.column_dimensions['B'].width = 15
        ws5.column_dimensions['C'].width = 15
        ws5.column_dimensions['D'].width = 20
        ws5.column_dimensions['E'].width = 12

    if bodega_data.get('por_proveedor'):
        ws6 = wb.create_sheet('Por Proveedor')
        headers6 = ['Proveedor', 'Existencia', 'Comprometida', 'Valor', 'Margen', 'Registros']
        for j, h in enumerate(headers6):
            cell = ws6.cell(row=1, column=j + 1, value=h)
            cell.font = header_font
            cell.fill = header_fill
        for i, (prov, data) in enumerate(bodega_data['por_proveedor'].items()):
            ws6.cell(row=i + 2, column=1, value=prov)
            ws6.cell(row=i + 2, column=2, value=data['existencia'])
            ws6.cell(row=i + 2, column=3, value=data['comprometida'])
            ws6.cell(row=i + 2, column=4, value=data['valor_total'])
            ws6.cell(row=i + 2, column=5, value=data['margen'])
            ws6.cell(row=i + 2, column=6, value=data['registros'])
        ws6.column_dimensions['A'].width = 40
        ws6.column_dimensions['B'].width = 15
        ws6.column_dimensions['C'].width = 15
        ws6.column_dimensions['D'].width = 20
        ws6.column_dimensions['E'].width = 10
        ws6.column_dimensions['F'].width = 12

    if bodega_data.get('top_items'):
        ws7 = wb.create_sheet('Top Items')
        headers7 = ['Item', 'Existencia', 'Comprometida', 'Valor', 'Precio Unit.', 'Registros']
        for j, h in enumerate(headers7):
            cell = ws7.cell(row=1, column=j + 1, value=h)
            cell.font = header_font
            cell.fill = header_fill
        for i, item in enumerate(bodega_data['top_items']):
            ws7.cell(row=i + 2, column=1, value=item['item'])
            ws7.cell(row=i + 2, column=2, value=item['existencia'])
            ws7.cell(row=i + 2, column=3, value=item['comprometida'])
            ws7.cell(row=i + 2, column=4, value=item['valor_total'])
            ws7.cell(row=i + 2, column=5, value=item['precio_unitario'])
            ws7.cell(row=i + 2, column=6, value=item['registros'])
        ws7.column_dimensions['A'].width = 50
        ws7.column_dimensions['B'].width = 15
        ws7.column_dimensions['C'].width = 15
        ws7.column_dimensions['D'].width = 20
        ws7.column_dimensions['E'].width = 15
        ws7.column_dimensions['F'].width = 12

    cache_key = f"report_{bodega_name.strip().upper()}"
    informe = current_data.get(cache_key, "")
    if informe:
        ws_report = wb.create_sheet('Informe de IA')
        lines = informe.split('\n')
        row = 1
        i = 0
        while i < len(lines):
            stripped = lines[i].strip()

            if stripped.startswith('# '):
                ws_report.cell(row=row, column=1, value=stripped[2:]).font = Font(name='Calibri', bold=True, size=16, color='4F46E5')
                ws_report.merge_cells(f'A{row}:F{row}')
                row += 1
            elif stripped.startswith('## '):
                ws_report.cell(row=row, column=1, value=stripped[3:]).font = Font(name='Calibri', bold=True, size=13, color='4F46E5')
                ws_report.merge_cells(f'A{row}:F{row}')
                row += 1
            elif stripped.startswith('### '):
                ws_report.cell(row=row, column=1, value=stripped[4:]).font = Font(name='Calibri', bold=True, size=11, color='334155')
                ws_report.merge_cells(f'A{row}:F{row}')
                row += 1

            elif stripped.startswith('|') and '|' in stripped[1:]:
                table_rows = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    row_line = lines[i].strip()
                    cells = [c.strip() for c in row_line.split('|')[1:-1]]
                    is_sep = all(c.replace('-', '').replace(':', '').strip() == '' for c in cells)
                    if not is_sep:
                        table_rows.append(cells)
                    i += 1
                i -= 1

                if table_rows:
                    num_cols = max(len(r) for r in table_rows)
                    for ri, row_data in enumerate(table_rows):
                        for ci, cell_text in enumerate(row_data):
                            if ci < num_cols:
                                cell = ws_report.cell(row=row, column=ci + 1, value=cell_text)
                                if ri == 0:
                                    cell.font = header_font
                                    cell.fill = header_fill
                                else:
                                    cell.font = Font(name='Calibri', size=10)
                                cell.border = thin_border
                                cell.alignment = Alignment(wrap_text=True)
                        row += 1
                    row += 1

            elif stripped.startswith('- '):
                ws_report.cell(row=row, column=1, value='  *  ' + stripped[2:]).font = Font(name='Calibri', size=10)
                row += 1
            elif stripped:
                ws_report.cell(row=row, column=1, value=stripped).font = Font(name='Calibri', size=10)
                row += 1
            i += 1

        ws_report.column_dimensions['A'].width = 30
        ws_report.column_dimensions['B'].width = 20
        ws_report.column_dimensions['C'].width = 20
        ws_report.column_dimensions['D'].width = 20
        ws_report.column_dimensions['E'].width = 20
        ws_report.column_dimensions['F'].width = 20

    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)

    safe_name = bodega_name.replace(' ', '_').replace('.', '').replace('/', '_')
    return StreamingResponse(
        buffer,
        media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': f'attachment; filename="Informe_{safe_name}.xlsx"'}
    )


VENDOR_EMAILS = {
    "INES SANCHEZ": "ines.sanchez@interdoors.com.co",
    "ELIANA": "eliana.gonzalez@interdoors.com.co",
    "KAROLIN": "karolin.gonzalez@interdoors.com.co",
    "MATEO POSADA": "mateo.posada@interdoors.com.co",
    "LEONARDO": "asesor3@interdoors.com.co",
    "YUDY CARRASQUILLA": "yudy.carrasquilla@interdoors.com.co",
    "LAURA OCHOA": "laura.ochoa@interdoors.com.co",
}


@app.get("/vendors")
async def list_vendors():
    raw_df = current_data.get("raw_df")
    if raw_df is None:
        raise HTTPException(status_code=400, detail="No hay datos cargados.")
    vendors = get_vendors(raw_df)
    for v in vendors:
        v["email"] = VENDOR_EMAILS.get(v["nombre"].upper(), "")
    return {"vendors": vendors}


def _generate_vendor_word_doc(vendor_name: str, items: list[dict]) -> io.BytesIO:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    title = doc.add_heading('', level=0)
    run = title.add_run(f'Informe de Inventario - {vendor_name}')
    run.font.color.rgb = RGBColor(79, 70, 229)
    run.font.size = Pt(22)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f'Generado el {datetime.now().strftime("%d/%m/%Y")}')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph()

    total_existencia = sum(i.get('existencia', 0) or 0 for i in items)
    total_comprometida = sum(i.get('cant_comprometida', 0) or 0 for i in items)
    total_disponible = sum(i.get('cant_disponible', 0) or 0 for i in items)
    total_valor = sum(i.get('valor_total', 0) or 0 for i in items)
    refs_unicas = len(set(i.get('referencia', '') for i in items if i.get('referencia')))
    total_registros = len(items)

    doc.add_heading('Resumen General', level=1)
    table = doc.add_table(rows=7, cols=2, style='Light Shading Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    metrics = [
        ('Total Registros', f"{total_registros:,}"),
        ('Referencias Unicas', f"{refs_unicas:,}"),
        ('Existencia Total', f"{total_existencia:,.0f}"),
        ('Cant. Comprometida', f"{total_comprometida:,.0f}"),
        ('Cant. Disponible', f"{total_disponible:,.0f}"),
        ('Valor Total', f"${total_valor:,.0f}"),
        ('Compromiso (%)', f"{(total_comprometida / total_existencia * 100) if total_existencia > 0 else 0:.1f}%"),
    ]
    for i, (label, value) in enumerate(metrics):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        for cell in table.rows[i].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    bodegas_map = {}
    for item in items:
        b = item.get('desc_bodega') or item.get('bodega') or 'Sin bodega'
        if b not in bodegas_map:
            bodegas_map[b] = []
        bodegas_map[b].append(item)

    if bodegas_map:
        doc.add_paragraph()
        doc.add_heading('Inventario por Bodega', level=1)
        for bodega_name, bodega_items in sorted(bodegas_map.items()):
            doc.add_heading(bodega_name, level=2)
            b_existencia = sum(i.get('existencia', 0) or 0 for i in bodega_items)
            b_valor = sum(i.get('valor_total', 0) or 0 for i in bodega_items)
            p = doc.add_paragraph()
            p.add_run(f'Registros: {len(bodega_items)} | ')
            p.add_run(f'Existencia: {b_existencia:,.0f} | ')
            p.add_run(f'Valor: ${b_valor:,.0f}')

            show_items = bodega_items[:20]
            if show_items:
                items_table = doc.add_table(rows=len(show_items) + 1, cols=5, style='Light List Accent 1')
                header_row = items_table.rows[0]
                for idx, h in enumerate(['Referencia', 'Descripcion', 'Existencia', 'Valor', 'Precio']):
                    header_row.cells[idx].text = h
                    for run in header_row.cells[idx].paragraphs[0].runs:
                        run.bold = True
                for j, item in enumerate(show_items):
                    row = items_table.rows[j + 1]
                    row.cells[0].text = str(item.get('referencia', ''))
                    row.cells[1].text = str(item.get('desc_item', ''))[:40]
                    row.cells[2].text = f"{(item.get('existencia', 0) or 0):,.0f}"
                    row.cells[3].text = f"${(item.get('valor_total', 0) or 0):,.0f}"
                    row.cells[4].text = f"${(item.get('precio_unitario', 0) or 0):,.0f}"
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(9)
            if len(bodega_items) > 20:
                doc.add_paragraph(f'... y {len(bodega_items) - 20} registros mas en esta bodega.')

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


@app.post("/send-emails")
async def send_emails(body: dict):
    raw_df = current_data.get("raw_df")
    if raw_df is None:
        raise HTTPException(status_code=400, detail="No hay datos cargados.")

    vendor_names = body.get("vendors", [])
    if not vendor_names:
        raise HTTPException(status_code=400, detail="Selecciona al menos un asesor.")

    smtp_user = os.environ.get("SMTP_USER", "")
    smtp_pass = os.environ.get("SMTP_PASS", "")
    if not smtp_user or not smtp_pass:
        raise HTTPException(status_code=500, detail="Credenciales SMTP no configuradas. Configura SMTP_USER y SMTP_PASS en .env")

    sent = 0
    errors = []

    for vendor_name in vendor_names:
        email = VENDOR_EMAILS.get(vendor_name.upper(), "")
        if not email:
            errors.append(f"{vendor_name}: email no encontrado en la lista de configuracion")
            continue

        items = get_vendor_items(raw_df, vendor_name)
        if not items:
            errors.append(f"{vendor_name}: no se encontraron registros con este nombre")
            continue

        doc_buffer = _generate_vendor_word_doc(vendor_name, items)

        total_existencia = sum(i.get('existencia', 0) or 0 for i in items)
        total_valor = sum(i.get('valor_total', 0) or 0 for i in items)
        refs = len(set(i.get('referencia', '') for i in items if i.get('referencia')))

        subject = f"Informe de Inventario - {vendor_name} - {datetime.now().strftime('%d/%m/%Y')}"
        body_text = (
            f"Hola {vendor_name},\n\n"
            f"Adjunto encontras tu informe de inventario actualizado.\n\n"
            f"Resumen:\n"
            f"- Registros totales: {len(items):,}\n"
            f"- Referencias unicas: {refs:,}\n"
            f"- Existencia total: {total_existencia:,.0f}\n"
            f"- Valor total: ${total_valor:,.0f}\n\n"
            f"Saludos,\n"
            f"Comercial - Interdoors"
        )

        try:
            import smtplib
            from email.mime.multipart import MIMEMultipart
            from email.mime.text import MIMEText
            from email.mime.base import MIMEBase
            from email import encoders

            msg = MIMEMultipart()
            msg['From'] = smtp_user
            msg['To'] = email
            msg['Subject'] = subject
            msg.attach(MIMEText(body_text, 'plain'))

            doc_buffer.seek(0)
            part = MIMEBase('application', 'vnd.openxmlformats-officedocument.wordprocessingml.document')
            part.set_payload(doc_buffer.read())
            encoders.encode_base64(part)
            safe_name = vendor_name.replace(' ', '_')
            part.add_header('Content-Disposition', f'attachment; filename="Informe_{safe_name}.docx"')
            msg.attach(part)

            with smtplib.SMTP('smtp.gmail.com', 587) as server:
                server.starttls()
                server.login(smtp_user, smtp_pass)
                server.send_message(msg)

            sent += 1
        except Exception as e:
            errors.append(f"{vendor_name}: {str(e)}")

    return {"sent": sent, "errors": errors, "total": len(vendor_names)}


if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", 8001))
    uvicorn.run(app, host="0.0.0.0", port=port)
